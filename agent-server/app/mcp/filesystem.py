"""filesystem MCP 래퍼 — 파일 Read / Create / Update / Backup / Delete
지원 포맷: 텍스트 (utf-8), Word (.docx), Excel (.xlsx)
"""
import os
import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path

from app.core.config import settings
from app.core.exceptions import (
    AgentFileExistsError,
    AgentFileNotFoundError,
    PathNotAllowedError,
)

# Office 포맷 지원 (ImportError 시 graceful degradation)
try:
    import docx as _docx          # python-docx
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    import openpyxl as _openpyxl  # openpyxl
    _XLSX_OK = True
except ImportError:
    _XLSX_OK = False


# ── 경로 검증 ──────────────────────────────────────────────────────────────────

def _nfc(p: Path) -> str:
    """macOS APFS는 경로를 NFD로 반환하고 Python str은 NFC — 양쪽을 NFC로 통일."""
    return unicodedata.normalize("NFC", str(p))


def _assert_allowed(path: str) -> Path:
    p = Path(path.strip()).resolve()
    allowed = [Path(d.strip()).resolve() for d in settings.ALLOWED_DIRECTORIES if d.strip()]
    if not allowed:
        return p  # 허용 경로 미설정 시 전체 허용 (개발 환경)
    p_str = _nfc(p)
    for a in allowed:
        a_str = _nfc(a)
        # os.sep 경계 검사 — /foo/bar2 가 /foo/bar 허용 시 통과되는 오류 방지
        if p_str == a_str or p_str.startswith(a_str + os.sep):
            return p
    raise PathNotAllowedError(path)


def _full_path(parent: Path, name: str) -> str:
    """디렉토리 항목의 절대 경로 문자열을 반환."""
    return str(parent / name)


# ── 디렉토리 조회 ─────────────────────────────────────────────────────────────

def list_directory(path: str) -> dict:
    p = _assert_allowed(path)
    if not p.exists():
        raise AgentFileNotFoundError(path)
    if not p.is_dir():
        raise AgentFileNotFoundError(path)

    items = []
    for item in sorted(p.iterdir(), key=lambda x: (x.is_file(), x.name.lower())):
        stat = item.stat()
        items.append({
            "name":        item.name,
            "path":        str(item),          # 절대 경로 (프론트엔드 탐색 / AI 도구 호출용)
            "type":        "directory" if item.is_dir() else "file",
            "size":        stat.st_size if item.is_file() else 0,
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        })
    return {"path": str(p), "items": items}


# ── 파일 읽기 ─────────────────────────────────────────────────────────────────

def read_file(path: str) -> dict:
    p = _assert_allowed(path)
    if not p.exists():
        raise AgentFileNotFoundError(path)

    ext = p.suffix.lower()

    if ext == ".docx":
        content = _read_docx(p)
        return {"path": str(p), "content": content, "size": p.stat().st_size, "encoding": "docx"}

    if ext == ".xlsx":
        content = _read_xlsx(p)
        return {"path": str(p), "content": content, "size": p.stat().st_size, "encoding": "xlsx"}

    # 기본: UTF-8 텍스트
    content = p.read_text(encoding="utf-8", errors="replace")
    return {"path": str(p), "content": content, "size": p.stat().st_size, "encoding": "utf-8"}


def _read_docx(p: Path) -> str:
    if not _DOCX_OK:
        return "[python-docx 미설치 — Word 파일을 읽으려면 pip install python-docx 를 실행하세요]"
    doc = _docx.Document(str(p))
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # 표 내용도 포함
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _read_xlsx(p: Path) -> str:
    if not _XLSX_OK:
        return "[openpyxl 미설치 — Excel 파일을 읽으려면 pip install openpyxl 을 실행하세요]"
    wb = _openpyxl.load_workbook(str(p), data_only=True)
    sections = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # None 셀을 빈 문자열로
            rows.append("\t".join("" if v is None else str(v) for v in row))
        sections.append(f"=== {sheet_name} ===\n" + "\n".join(rows))
    return "\n\n".join(sections)


# ── 백업 ────────────────────────────────────────────────────────────────────

def backup_file(src_path: str, dest_path: str = "") -> dict:
    src = _assert_allowed(src_path)
    if not src.exists():
        raise AgentFileNotFoundError(src_path)

    if dest_path:
        dst = _assert_allowed(dest_path)
    else:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        dst = src.parent / f"{src.stem}.backup_{ts}{src.suffix}"

    shutil.copy2(src, dst)
    return {"src_path": str(src), "backup_path": str(dst)}


# ── 파일 생성 ─────────────────────────────────────────────────────────────────

def create_file(path: str, content: str, overwrite: bool = False) -> dict:
    p = _assert_allowed(path)
    if p.exists() and not overwrite:
        raise AgentFileExistsError(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    ext = p.suffix.lower()
    if ext == ".docx":
        _write_docx(p, content)
    elif ext == ".xlsx":
        _write_xlsx(p, content)
    else:
        p.write_text(content, encoding="utf-8")

    return {"path": str(p), "size": p.stat().st_size}


# ── 파일 수정 ─────────────────────────────────────────────────────────────────

def update_file(path: str, content: str) -> dict:
    p = _assert_allowed(path)
    if not p.exists():
        raise AgentFileNotFoundError(path)

    ext = p.suffix.lower()
    if ext == ".docx":
        _write_docx(p, content)
    elif ext == ".xlsx":
        _write_xlsx(p, content)
    else:
        p.write_text(content, encoding="utf-8")

    return {"path": str(p), "size": p.stat().st_size}


def _write_docx(p: Path, content: str) -> None:
    """마크다운 문법을 파싱하여 Word 서식으로 변환 후 저장.

    지원 문법:
      # ~ ####  → Heading 1~4
      - / *     → 글머리 기호 목록 (List Bullet)
      1. 2. …   → 번호 목록 (List Number)
      **text**  → 굵게
      *text*    → 기울임
      ***text***→ 굵게+기울임
      ---       → 단락 구분 (빈 줄)
      빈 줄      → 단락 간격
    """
    if not _DOCX_OK:
        raise RuntimeError("python-docx 미설치 — pip install python-docx")

    doc = _docx.Document()

    for line in content.splitlines():
        stripped = line.rstrip()

        if stripped.startswith("#### "):
            _docx_paragraph(doc, stripped[5:], "Heading 4")
        elif stripped.startswith("### "):
            _docx_paragraph(doc, stripped[4:], "Heading 3")
        elif stripped.startswith("## "):
            _docx_paragraph(doc, stripped[3:], "Heading 2")
        elif stripped.startswith("# "):
            _docx_paragraph(doc, stripped[2:], "Heading 1")
        elif re.match(r"^[-*] ", stripped):
            _docx_paragraph(doc, stripped[2:], "List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            _docx_paragraph(doc, text, "List Number")
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph()   # 수평선 → 빈 단락으로 대체
        elif not stripped:
            doc.add_paragraph()
        else:
            _docx_paragraph(doc, stripped, "Normal")

    doc.save(str(p))


def _docx_paragraph(doc, text: str, style: str):
    """스타일이 적용된 단락에 인라인 마크다운 서식(**bold**, *italic*)을 추가."""
    para = doc.add_paragraph(style=style)
    _docx_inline(para, text)
    return para


def _docx_inline(para, text: str) -> None:
    """***bold+italic***, **bold**, *italic* 인라인 서식을 Run으로 분해하여 적용."""
    pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):          # ***bold+italic***
            run = para.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2):        # **bold**
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):        # *italic*
            run = para.add_run(m.group(3))
            run.italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _write_xlsx(p: Path, content: str) -> None:
    """탭 또는 쉼표로 구분된 텍스트를 xlsx로 저장.

    지원 기능:
      - 첫 행 굵게 + 가운데 정렬 (헤더로 간주)
      - 열 너비 자동 조정
      - '=== 시트명 ===' 줄로 다중 시트 구분
    """
    if not _XLSX_OK:
        raise RuntimeError("openpyxl 미설치 — pip install openpyxl")

    from openpyxl.styles import Alignment, Font

    wb = _openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    is_first_row = True

    for line in content.splitlines():
        stripped = line.strip()

        # === 시트명 === → 새 시트 전환
        m = re.match(r"^===\s*(.+?)\s*===$", stripped)
        if m:
            sheet_title = m.group(1)[:31]
            if ws.max_row == 0:
                ws.title = sheet_title
            else:
                ws = wb.create_sheet(title=sheet_title)
            is_first_row = True
            continue

        if not stripped:
            continue

        if "\t" in line:
            cells = line.split("\t")
        elif "," in stripped:
            cells = [c.strip() for c in stripped.split(",")]
        else:
            cells = [stripped]

        row_idx = ws.max_row + 1
        for col_idx, val in enumerate(cells, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            if is_first_row:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")

        is_first_row = False

    # 열 너비 자동 조정 (최대 50)
    for sheet in wb.worksheets:
        for col in sheet.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=0)
            sheet.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    wb.save(str(p))


# ── 파일 삭제 ─────────────────────────────────────────────────────────────────

def delete_file(path: str) -> dict:
    """파일 삭제 전 자동 백업을 생성한다. 디렉토리는 삭제 불가."""
    p = _assert_allowed(path)
    if not p.exists():
        raise AgentFileNotFoundError(path)
    if p.is_dir():
        raise AgentFileNotFoundError(f"{path} (디렉토리는 삭제할 수 없습니다)")

    # 자동 백업
    backup = backup_file(path)
    p.unlink()
    return {"deleted_path": str(p), "backup_path": backup["backup_path"]}
